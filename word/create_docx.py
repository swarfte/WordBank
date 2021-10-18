'''
Author: Chau Lap Tou
Date: 2021-08-31 16:33:37
LastEditors: Chau Lap Tou
LastEditTime: 2021-08-31 17:25:49
python_exe: pyinstaller -F -w file_name.py -p C:/python/lib/site-packages 
java_class: javac -encoding utf-8 file_name.java
java_jar: jar -cvmf manifest.txt name.jar *.class
GithubName: Swarfte
'''
import docx


def build_docx_file_all(path, value, row=1, col=5):  # *108+1行
    doc = docx.Document()  # *創建一個doc文件
    table = doc.add_table(rows=row, cols=col)  # *創建一個表格
    hdr_cells = table.rows[0].cells  # *設置為頭頂的行
    hdr_cells[0].text = "Chinese"  # *cell表示列
    hdr_cells[1].text = "adjective"
    hdr_cells[2].text = "noun"
    hdr_cells[3].text = "verb"
    hdr_cells[4].text = 'adverb'

    data = value

    for a, b, c, d, e in data:
        row_cells = table.add_row().cells  # *用作表示表格的列
        row_cells[0].text = a  # *第一列
        row_cells[1].text = b  # *第二列
        row_cells[2].text = c  # *第三列
        row_cells[3].text = d  # *第四列
        row_cells[4].text = e  # *第五列

    doc.save(path)  # *保存文件


def build_chinese_and_english(path, value, row=1, col=2):
    doc = docx.Document()  # *創建一個doc文件
    table = doc.add_table(rows=row, cols=col)  # *創建一個表格
    hdr_cells = table.rows[0].cells  # *設置為頭頂的行
    hdr_cells[0].text = "English"  # *cell表示列
    hdr_cells[1].text = "中文"
    data = value
    for x in range(len(data[0])):
        row_cells = table.add_row().cells  # *用作表示表格的列
        row_cells[0].text = data[0][x]  # *第一列
        row_cells[1].text = data[1][x]  # *第二列
    doc.save(path)  # *保存文件


def build_any_file(path, value, row, col, col_name): #適用於任意檔案
    doc = docx.Document()  # *創建一個doc文件
    table = doc.add_table(rows=row, cols=col)
    hdr_cells = table.rows[0].cells
    for x in range(col):  # 設置每一欄的名稱
        hdr_cells[x].text = col_name[x]

    max_length = 0
    for x in value:
        if len(x) > max_length:
            max_length = len(x)

    for x in range(max_length):  # 總運行次數
        row_cells = table.add_row().cells
        for y in range(value):  # 每一列在對應行的值
            try:
                row_cells[y].text = value[y][x]
            except Exception as ex: #空指針的情況
                row_cells[y].text = ""

    doc.save(path)


if __name__ == "__main__":
    pass
